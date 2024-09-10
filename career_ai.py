import streamlit as st
from deep_translator import GoogleTranslator
import docx2txt
from PyPDF2 import PdfReader
import openai
from io import BytesIO
from fpdf import FPDF
from docx import Document
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Set OpenAI API key
openai.api_key = OPENAI_API_KEY

# Function to add CSS styling
def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

# Add your custom CSS file
local_css("style.css")

# Function to extract text from CV
def extract_text_from_cv(file):
    text = ""
    if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = docx2txt.process(file)
    elif file.type == "application/pdf":
        try:
            reader = PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
        except Exception as e:
            st.error(f"Error extracting text from PDF: {e}")
    return text

# Function to recommend jobs based on career
def recommend_jobs_based_on_career(career):
    career_jobs = {
        "Software Developer": ["Full-Stack Developer", "Backend Engineer", "Frontend Developer"],
        "Data Scientist": ["Machine Learning Engineer", "Data Analyst", "AI Researcher"],
        "Nurse": ["Registered Nurse", "Pediatric Nurse", "Clinical Nurse"],
        "Teacher": ["Math Teacher", "Science Teacher", "English Teacher"],
        "Other": ["Freelance Consultant", "Entrepreneur", "Researcher"]
    }
    return career_jobs.get(career, [])

# Function to refine job recommendations based on CV content
def refine_jobs_based_on_cv(cv_text, initial_jobs):
    skills_mentioned = ["Python", "Django", "Data Analysis", "Machine Learning"]
    refined_jobs = []
    for job in initial_jobs:
        for skill in skills_mentioned:
            if skill.lower() in cv_text.lower():
                refined_jobs.append(job)
                break
    return refined_jobs if refined_jobs else initial_jobs

# Hybrid function to recommend jobs based on career and CV
def recommend_jobs(career, cv_text):
    initial_jobs = recommend_jobs_based_on_career(career)
    final_jobs = refine_jobs_based_on_cv(cv_text, initial_jobs)
    return final_jobs

# Function to generate a career overview based on career selection
def generate_career_overview(career):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert career advisor."},
                {"role": "user", "content": f"Provide a professional overview for {career}"}
            ],
            max_tokens=700
        )
        return response.choices[0].message['content'].strip()
    except openai.error.AuthenticationError:
        st.error("Invalid OpenAI API key. Please check your API key.")
    except Exception as e:
        st.error(f"Error generating career overview: {e}")

# Function to refine CV with direct corrections
def generate_refined_cv(cv_text, career):
    prompt = f"""
    You are an expert career advisor. A user has uploaded a CV for the career of {career}.
    Please directly correct and improve the CV to increase the chances of getting a job in {career}.
    Do not explain the changes; just provide the corrected CV.

    Original CV:
    {cv_text}
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert career advisor."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=700
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        st.error(f"Error generating refined CV: {e}")
        return "Error generating refined CV."

# Function for generating direct mock interview responses
def generate_mock_interview_response(question):
    prompt = f"""
    You are a professional interviewer. A candidate has asked the following question: '{question}'.
    Please provide a straightforward, professional response to this question, focused on the specific topic.
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional interviewer."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=300
        )
        return response.choices[0].message['content'].strip()
    except openai.error.AuthenticationError:
        st.error("Invalid OpenAI API key. Please check your API key.")
    except openai.error.RateLimitError:
        st.error("Rate limit exceeded. Please try again later.")
    except Exception as e:
        st.error(f"Error generating mock interview response: {e}")

# Function to generate PDF file
def create_pdf(cv_content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, cv_content)
    
    buffer = BytesIO()
    pdf_output = pdf.output(dest='S').encode('latin1')  # Encode the output to handle binary data
    buffer.write(pdf_output)
    buffer.seek(0)
    
    return buffer


# Function to generate Word file (DOCX)
def create_word(cv_content):
    doc = Document()
    doc.add_heading('Refined CV', 0)
    for line in cv_content.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to translate text
def translate_text(text, lang_code):
    try:
        translation = GoogleTranslator(source='auto', target=lang_code).translate(text)
        return translation
    except Exception as e:
        st.error(f"Translation Error: {e}")
        return text

# Define all texts used in the system
texts = {
    "title": "AI-Powered Career Guidance System",
    "interview_prep_header": "Interview Preparation",
    "welcome_message": "Welcome to the Interview Preparation section! Here, you'll find common interview questions, tips, and even a mock interview simulator to help you practice.",
    "common_questions_header": "Common Interview Questions",
    "common_questions": [
        "Tell me about yourself.",
        "Why do you want to work here?",
        "What are your strengths and weaknesses?",
        "Describe a challenge you faced and how you handled it.",
        "Where do you see yourself in five years?"
    ],
    "interview_tips_header": "Interview Tips",
    "interview_tips": [
        "Research the company and role before the interview.",
        "Practice common interview questions and answers.",
        "Dress appropriately for the interview.",
        "Be punctual and respectful.",
        "Follow up with a thank you note after the interview."
    ],
    "mock_interview_header": "Mock Interview Simulator",
    "career_selection_header": "Select or Enter Your Career",
    "upload_cv_header": "Upload your CV",
    "multilingual_support_header": "Multilingual Support"
}

# Streamlit app
def main():
    st.subheader(texts["multilingual_support_header"])
    lang = st.selectbox("Select Language", ["English", "Afrikaans", "Zulu", "Xhosa", "Sepedi", "Setswana", "Sesotho", "Xitsonga", "SiSwati", "Tshivenda"])
    lang_code = {
        "English": "en",
        "Afrikaans": "af",
        "Zulu": "zu",
        "Xhosa": "xh",
        "Sepedi": "nso",
        "Setswana": "tn",
        "Sesotho": "st",
        "Xitsonga": "ts",
        "SiSwati": "ss",
        "Tshivenda": "ve"
    }.get(lang, "en")

    translated_texts = {key: translate_text(value, lang_code) if isinstance(value, str) else [translate_text(q, lang_code) for q in value] for key, value in texts.items()}

    st.title(translated_texts["title"])

    st.header(translated_texts["interview_prep_header"])
    st.markdown(translated_texts["welcome_message"])

    st.subheader(translated_texts["common_questions_header"])
    st.write(translated_texts["common_questions"])

    st.subheader(translated_texts["interview_tips_header"])
    st.write(translated_texts["interview_tips"])

    st.subheader(translated_texts["mock_interview_header"])
    user_question = st.text_input("Ask an interview question:")
    if st.button("Simulate Interview"):
        if user_question:
            with st.spinner("Generating interview response..."):
                ai_response = generate_mock_interview_response(user_question)
            st.write("AI's Response:", ai_response)
        else:
            st.write("Please enter a question.")
    
    st.subheader(translated_texts["career_selection_header"])
    careers = ["Software Developer", "Data Scientist", "Nurse", "Teacher", "Other"]
    selected_career = st.selectbox("Select Career", careers)
    st.write("Selected Career:", selected_career)

    uploaded_file = st.file_uploader(translated_texts["upload_cv_header"], type=["pdf", "docx"])
    if uploaded_file:
        cv_text = extract_text_from_cv(uploaded_file)
        if cv_text:
            st.text_area("CV Content", cv_text, height=200)

            if st.button("Generate Career Overview"):
                with st.spinner("Generating career overview..."):
                    overview = generate_career_overview(selected_career)
                st.write("Career Overview:", overview)

            if st.button("Generate Refined CV"):
                with st.spinner("Generating refined CV..."):
                    refined_cv = generate_refined_cv(cv_text, selected_career)
                st.write("Refined CV:", refined_cv)
                
                # Provide options to download refined CV
                pdf_buffer = create_pdf(refined_cv)
                st.download_button("Download PDF", pdf_buffer, file_name="Refined_CV.pdf", mime="application/pdf")
                
                word_buffer = create_word(refined_cv)
                st.download_button("Download Word Document", word_buffer, file_name="Refined_CV.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            if st.button("Recommend Jobs"):
                with st.spinner("Recommending jobs..."):
                    recommended_jobs = recommend_jobs(selected_career, cv_text)
                st.write("Recommended Jobs:", recommended_jobs)

if __name__ == "__main__":
    main()
